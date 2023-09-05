from typing import Optional
import docx as d
import requests
from tkinter import filedialog as fd

# Print iterations progress
def printProgressBar (iteration, total, prefix = '', suffix = '', decimals = 1, length = 100, fill = '█', printEnd = "\r"):
    """
    Call in a loop to create terminal progress bar
    @params:
        iteration   - Required  : current iteration (Int)
        total       - Required  : total iterations (Int)
        prefix      - Optional  : prefix string (Str)
        suffix      - Optional  : suffix string (Str)
        decimals    - Optional  : positive number of decimals in percent complete (Int)
        length      - Optional  : character length of bar (Int)
        fill        - Optional  : bar fill character (Str)
        printEnd    - Optional  : end character (e.g. "\r", "\r\n") (Str)
    """
    percent = ("{0:." + str(decimals) + "f}").format(100 * (iteration / float(total)))
    filledLength = int(length * iteration // total)
    bar = fill * filledLength + '-' * (length - filledLength)
    print(f'\r{prefix} |{bar}| {percent}% {suffix}'[:140], end = printEnd)
    # Print New Line on Complete
    if iteration == total: 
        print()

PADDING = " "

def improveKanjiList(kanjiList:str, out:Optional[str] = None) -> None:
    """Searches a Docx file, gets the first table's second column
    Assuming that it contains the kanji. Then uses get def to place 
    definitions in the 4th column of the table then saves the new kanji list

    Args:
        kanjiList (str): The path of the docx file to find
        out (Optional[str]): The path of the docx file to save to.
    """
    if not kanjiList:
        return
    doc = d.Document(kanjiList)
    m_table = doc.tables[0]
    kanjis = m_table.column_cells(3)
    translations = [len(c.paragraphs) for c in kanjis[1:]]
    total = sum(translations)
    print(f"Starting Search for {total} Kanji compounds...")
    printProgressBar(0,len(kanjis[1:]),"","",1,10)
    last = ""
    for i,c in enumerate(kanjis[1:]):
        m_table.cell(i+1,4).text = ""
        for j,x in enumerate(c.paragraphs):
            meaning = getDef(x.text)
            text = f"{i+1}.{j+1}:{x.text} - {meaning}"
            mult = len(last) - len(text) if len(last) > len(text) else 0
            printProgressBar(sum(translations[:i]) + j + 1,total,"",f"{text}{PADDING*mult}",1,10)
            m_table.cell(i+1,4).text += "\n" + f"{j+1}-{meaning}" if j >= 1 else f"{j+1}-{meaning}"
            last = text
    if out :
        doc.save(out)
    else:
        doc.save(kanjiList[:kanjiList.rfind("/") + 1] + "new " + kanjiList[kanjiList.rfind("/") + 1:] if "/" in kanjiList else "new " + kanjiList )


def getDef(kanji:str) -> str:
    """Gets the definition of a kanji or kanji compound from https://jisho.org

    Args:
        kanji (str): The str containing the kanji definition to find

    Returns:
        str: The string definition of the kanji
    """
    if len(kanji) == 1:
        with requests.get(f"https://jisho.org/search/{kanji}%20%23kanji") as out:
            start = out.text.find('<div class="kanji-details__main-meanings">',out.text.rfind('</head>'))
            return out.text[start+len('<div class="kanji-details__main-meanings">') + 7:out.text.find('</div>',start) - 5]
    with requests.get(f"https://jisho.org/search/{kanji}") as out:
        start = out.text.find('<span class="meaning-meaning">',out.text.rfind('</head>'))
        return out.text[start+len('<span class="meaning-meaning">'):out.text.find('</span>',start)]

def getFurigana(kanji: str) -> str:
    pass

if __name__ == "__main__":
    improveKanjiList(fd.askopenfilename(
        title="File For translate",
        initialdir="./",
        filetypes = (
            ('Word files', '*.docx'),
        )
    ))
